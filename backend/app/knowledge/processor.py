"""Document text extraction and chunking driven by the canvas Chunking node config."""

from __future__ import annotations

import os
import shutil
from pathlib import Path
from typing import Any, Callable

from ..canvas.models import CanvasFlow
from .models import KnowledgeChunk
from .tokenizers import chars_per_token, get_token_counter


# Default fallback chunking parameters when the canvas flow has no chunking node.
DEFAULT_CHUNK_SIZE = 700
DEFAULT_OVERLAP = 120
DEFAULT_STRATEGY = "recursive"
# Used only when the embedding tokenizer is unavailable.
FALLBACK_CHARS_PER_TOKEN = 4.0


def _configure_tesseract_path() -> None:
    """Point pytesseract at the Tesseract binary if it isn't on PATH (Windows)."""
    try:
        import pytesseract  # type: ignore
    except ImportError:
        return
    if shutil.which("tesseract"):
        return
    for candidate in (
        os.environ.get("TESSERACT_CMD"),
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
        "/opt/homebrew/bin/tesseract",
    ):
        if candidate and os.path.isfile(candidate):
            pytesseract.pytesseract.tesseract_cmd = candidate
            break


_configure_tesseract_path()


class KnowledgeProcessor:
    """Extract text from common document types and split it into chunks."""

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    def extract_text(self, file_path: Path, content_type: str = "") -> tuple[str, int | None]:
        """Return (full_text, page_count_or_none) from a single file on disk."""
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".pdf":
                return self._extract_pdf(file_path)
            if suffix in {".docx"}:
                return self._extract_docx(file_path), None
            if suffix in {".txt", ".md", ".markdown", ".log", ".csv", ".json", ".html", ".htm"}:
                return self._extract_text_file(file_path), None
        except Exception as exc:  # noqa: BLE001 — surfaced to caller as error status
            raise RuntimeError(f"Failed to extract text from '{file_path.name}': {exc}") from exc

        # Unknown type – best-effort UTF-8 decode, otherwise empty.
        return self._extract_text_file(file_path, lossy=True), None

    def _extract_pdf(self, file_path: Path) -> tuple[str, int]:
        """Extract text from a PDF, falling back to OCR for scanned pages.

        Strategy per page:
          1. Try PyMuPDF (fast, accurate) if installed — else pypdf.
          2. If a page yields almost no text (likely scanned), render it to an
             image and run Tesseract OCR via pytesseract.
        """
        # Prefer PyMuPDF for extraction + rendering (single dependency for both).
        try:
            import pymupdf  # type: ignore
        except ImportError:  # noqa: BLE001
            pymupdf = None  # type: ignore

        if pymupdf is not None:
            doc = pymupdf.open(str(file_path))
            try:
                pages: list[str] = []
                ocr_attempted = 0
                ocr_unavailable = False
                for page in doc:
                    text = ""
                    try:
                        text = (page.get_text("text") or "").strip()
                    except Exception:  # noqa: BLE001
                        text = ""
                    if len(text) < 20:  # likely scanned / image-only page
                        ocr_attempted += 1
                        ocr_text, ocr_status = self._ocr_pdf_page(page)
                        if ocr_status == "unavailable":
                            ocr_unavailable = True
                        if ocr_text:
                            text = ocr_text
                    pages.append(text)
                joined = "\n\n".join(pages).strip()
                if not joined and ocr_attempted and ocr_unavailable:
                    raise RuntimeError(
                        "PDF appears to be scanned and OCR is required, but Tesseract is not installed. "
                        "Install it (Windows: https://github.com/UB-Mannheim/tesseract/wiki) and ensure "
                        "'tesseract' is on PATH, then re-upload."
                    )
                return joined, doc.page_count
            finally:
                doc.close()

        # Fallback: pypdf (no OCR — image-only PDFs will return empty text).
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                "PDF support requires the 'pymupdf' or 'pypdf' package. "
                "Install with: pip install pymupdf pytesseract"
            ) from exc

        reader = PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                pages.append("")
        return "\n\n".join(pages).strip(), len(reader.pages)

    def _ocr_pdf_page(self, page: Any) -> tuple[str, str]:
        """Render a PyMuPDF page to an image and run Tesseract OCR.

        Returns ``(text, status)`` where ``status`` is one of:
          - ``"ok"``       — OCR ran successfully (text may still be empty)
          - ``"unavailable"`` — pytesseract or Tesseract binary not installed
          - ``"error"``    — OCR attempted but raised
        """
        try:
            import pytesseract  # type: ignore
            from PIL import Image  # type: ignore
        except ImportError:
            return "", "unavailable"

        try:
            import pymupdf  # type: ignore

            matrix = pymupdf.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            mode = "RGB" if pix.n < 4 else "RGBA"
            image = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
            try:
                text = pytesseract.image_to_string(image, lang="hun+eng") or ""
            except pytesseract.TesseractError:
                text = pytesseract.image_to_string(image) or ""
            except pytesseract.TesseractNotFoundError:
                return "", "unavailable"
            return text.strip(), "ok"
        except Exception:  # noqa: BLE001
            return "", "error"

    def _extract_docx(self, file_path: Path) -> str:
        try:
            from docx import Document  # type: ignore  # python-docx
        except ImportError as exc:
            raise RuntimeError(
                "DOCX support requires the 'python-docx' package. Install it with: pip install python-docx"
            ) from exc

        document = Document(str(file_path))
        parts = [para.text for para in document.paragraphs if para.text]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts).strip()

    def _extract_text_file(self, file_path: Path, lossy: bool = False) -> str:
        encodings = ("utf-8", "utf-8-sig", "latin-1")
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        if lossy:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        raise RuntimeError(f"Could not decode '{file_path.name}' as text.")

    # ------------------------------------------------------------------
    # Chunking config resolution
    # ------------------------------------------------------------------

    def resolve_chunking_config(self, flow: CanvasFlow | None) -> dict[str, Any]:
        """Pick the first ``chunking`` node from the flow and return its config.

        Returns a dict with the canonical keys ``chunkSize``, ``overlap``,
        ``strategy``, plus any extra keys the user configured on the node
        (``separators``, ``keepSeparator``, ``lengthFunction``,
        ``minChunkChars``, ``stripWhitespace``).
        """
        config: dict[str, Any] = {
            "chunkSize": DEFAULT_CHUNK_SIZE,
            "overlap": DEFAULT_OVERLAP,
            "strategy": DEFAULT_STRATEGY,
            "separators": "\\n\\n,\\n,. , ,",
            "keepSeparator": True,
            "lengthFunction": "characters",
            "minChunkChars": 0,
            "stripWhitespace": True,
        }
        if flow is None:
            return config
        for node in flow.nodes:
            # Canvas v2 uses ``process-chunking`` as the template key; older
            # blueprints used the bare ``chunking`` literal — accept both.
            if node.template_key in {"process-chunking", "chunking"}:
                node_config = dict(node.config or {})
                if "chunkSize" in node_config:
                    config["chunkSize"] = int(node_config["chunkSize"])
                if "overlap" in node_config:
                    config["overlap"] = int(node_config["overlap"])
                if "strategy" in node_config:
                    config["strategy"] = str(node_config["strategy"])
                # Preserve any additional keys (e.g. separators) verbatim.
                for key, value in node_config.items():
                    config[key] = value
                break
        # Bind the upstream embedding model so the chunker can switch to the
        # right tokenizer (multilingual e5 ≈ 2 chars/token, OpenAI ≈ 4).
        for node in flow.nodes:
            if node.template_key in {"process-embedding", "embedding"}:
                emb_cfg = node.config or {}
                metadata = emb_cfg.get("metadata") or {}
                model_id = (
                    metadata.get("model_id")
                    or emb_cfg.get("model_id")
                    or ""
                )
                if model_id:
                    config["embeddingModelId"] = model_id
                break
        return config

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def chunk_text(self, doc_id: str, text: str, config: dict[str, Any]) -> list[KnowledgeChunk]:
        chunk_size = max(50, int(config.get("chunkSize", DEFAULT_CHUNK_SIZE)))
        overlap = max(0, min(int(config.get("overlap", DEFAULT_OVERLAP)), chunk_size - 1))
        strategy = str(config.get("strategy", DEFAULT_STRATEGY)).lower()
        keep_separator = bool(config.get("keepSeparator", True))
        min_chunk_chars = int(config.get("minChunkChars", 0) or 0)
        strip_whitespace = bool(config.get("stripWhitespace", True))
        length_function = str(config.get("lengthFunction", "characters")).lower()
        separators = self._coerce_separators(config.get("separators"))
        embedding_model_id = config.get("embeddingModelId") or (
            (config.get("embeddingProfile") or {}).get("modelId")
        )

        text = (text or "").strip() if strip_whitespace else (text or "")
        if not text:
            return []

        # Token-aware mode: scale chunk_size/overlap from the *real* embedding
        # tokenizer when possible. We sample the actual document so the
        # chars-per-token ratio reflects the input language (Hungarian SP
        # tokens are smaller than English BPE tokens, etc.).
        token_counter = None
        token_budget = chunk_size  # remembered in tokens for the safety pass
        if length_function in {"tokens", "token"}:
            ratio = chars_per_token(embedding_model_id, text[:4000])
            if ratio is None:
                ratio = FALLBACK_CHARS_PER_TOKEN
            else:
                token_counter = get_token_counter(embedding_model_id)
            chunk_size = max(50, int(chunk_size * ratio))
            overlap = max(0, int(overlap * ratio))

        if strategy in {"recursive", "langchain", "recursivecharacter"}:
            pieces = self._recursive_split(text, chunk_size, overlap, separators, keep_separator)
        elif strategy in {"sentence", "sentences"}:
            pieces = self._sentence_split(text, chunk_size, overlap)
        elif strategy in {"markdown", "md"}:
            pieces = self._recursive_split(
                text,
                chunk_size,
                overlap,
                ["\n## ", "\n# ", "\n### ", "\n\n", "\n", ". ", " ", ""],
                keep_separator,
            )
        else:
            pieces = self._fixed_window_split(text, chunk_size, overlap)

        chunks: list[KnowledgeChunk] = []
        for index, piece in enumerate(pieces):
            cleaned = piece.strip() if strip_whitespace else piece
            if not cleaned or len(cleaned) < min_chunk_chars:
                continue
            # Safety pass: when a real tokenizer is bound, hard-cap the chunk
            # at the requested token budget so embeddings never silently
            # truncate. We split on whitespace and re-emit until the budget
            # is satisfied.
            sub_pieces = [cleaned]
            if token_counter is not None:
                sub_pieces = self._enforce_token_budget(cleaned, token_counter, token_budget)
            for sub in sub_pieces:
                if not sub or len(sub) < min_chunk_chars:
                    continue
                token_count = token_counter(sub) if token_counter else max(1, len(sub) // 4)
                chunks.append(
                    KnowledgeChunk(
                        id=f"{doc_id}-c{len(chunks):04d}",
                        index=len(chunks),
                        text=sub,
                        char_count=len(sub),
                        token_estimate=token_count,
                    )
                )
        return chunks

    @staticmethod
    def _enforce_token_budget(
        text: str,
        counter: Callable[[str], int],
        budget: int,
    ) -> list[str]:
        """Greedily pack words into pieces that each fit within ``budget`` tokens.

        Used as a final safety net after the character-based splitter so the
        embedding model never receives an over-budget input — which Pinecone
        integrated inference would silently truncate, harming recall.
        """
        if budget <= 0:
            return [text]
        if counter(text) <= budget:
            return [text]
        words = text.split(" ")
        out: list[str] = []
        buf: list[str] = []
        for word in words:
            candidate = (" ".join(buf + [word])).strip()
            if not candidate:
                continue
            if counter(candidate) > budget and buf:
                out.append(" ".join(buf).strip())
                buf = [word]
            else:
                buf.append(word)
        if buf:
            out.append(" ".join(buf).strip())
        return [piece for piece in out if piece]

    @staticmethod
    def _coerce_separators(raw: Any) -> list[str]:
        """Translate the user-facing separator string/list into a clean list.

        Accepts either a comma-separated string (with literal ``\\n``/``\\t`` escapes)
        or a list of separators.
        """
        def decode(token: str) -> str:
            return (
                token.replace("\\n", "\n").replace("\\t", "\t").replace("\\r", "\r")
            )

        if isinstance(raw, list):
            cleaned = [decode(str(part)) for part in raw]
            return cleaned or ["\n\n", "\n", ". ", " ", ""]
        if isinstance(raw, str) and raw.strip():
            cleaned = [decode(part) for part in raw.split(",") if part != ""]
            # Always make sure the recursive split has a fallback empty separator.
            if "" not in cleaned:
                cleaned.append("")
            return cleaned
        return ["\n\n", "\n", ". ", " ", ""]

    # ------------------------------------------------------------------
    # Splitting strategies
    # ------------------------------------------------------------------

    def _fixed_window_split(self, text: str, chunk_size: int, overlap: int) -> list[str]:
        step = max(1, chunk_size - overlap)
        pieces: list[str] = []
        for start in range(0, len(text), step):
            piece = text[start : start + chunk_size]
            if not piece:
                break
            pieces.append(piece)
            if start + chunk_size >= len(text):
                break
        return pieces

    def _recursive_split(
        self,
        text: str,
        chunk_size: int,
        overlap: int,
        separators: list[str] | None = None,
        keep_separator: bool = True,
    ) -> list[str]:
        seps = separators if separators is not None else ["\n\n", "\n", ". ", " ", ""]
        return self._split_with_separators(text, seps, chunk_size, overlap, keep_separator)

    def _split_with_separators(
        self,
        text: str,
        separators: list[str],
        chunk_size: int,
        overlap: int,
        keep_separator: bool = True,
    ) -> list[str]:
        if len(text) <= chunk_size:
            return [text] if text.strip() else []

        sep = separators[0] if separators else ""
        if sep == "":
            return self._fixed_window_split(text, chunk_size, overlap)

        parts = text.split(sep)
        # Re-attach the separator to all but the last fragment so we don't lose context.
        if keep_separator:
            rebuilt = [part + sep for part in parts[:-1]] + ([parts[-1]] if parts else [])
        else:
            rebuilt = parts

        chunks: list[str] = []
        buffer = ""
        for part in rebuilt:
            if len(part) > chunk_size:
                if buffer:
                    chunks.append(buffer)
                    buffer = ""
                chunks.extend(
                    self._split_with_separators(part, separators[1:], chunk_size, overlap, keep_separator)
                )
                continue
            if len(buffer) + len(part) <= chunk_size:
                buffer += part
            else:
                if buffer:
                    chunks.append(buffer)
                buffer = part
        if buffer:
            chunks.append(buffer)

        # Apply a soft overlap by prefixing each chunk with the tail of the previous one.
        if overlap > 0 and len(chunks) > 1:
            with_overlap: list[str] = [chunks[0]]
            for previous, current in zip(chunks, chunks[1:]):
                tail = previous[-overlap:]
                with_overlap.append(tail + current)
            return with_overlap
        return chunks

    def _sentence_split(self, text: str, chunk_size: int, overlap: int) -> list[str]:
        import re

        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks: list[str] = []
        buffer = ""
        for sentence in sentences:
            if not sentence.strip():
                continue
            if len(buffer) + len(sentence) + 1 <= chunk_size:
                buffer = f"{buffer} {sentence}".strip()
            else:
                if buffer:
                    chunks.append(buffer)
                if len(sentence) > chunk_size:
                    chunks.extend(self._fixed_window_split(sentence, chunk_size, overlap))
                    buffer = ""
                else:
                    buffer = sentence
        if buffer:
            chunks.append(buffer)
        return chunks
